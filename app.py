from datetime import datetime
import io
import random
import pandas as pd
import streamlit as st
import sqlite3
from database import *


# --- НАСТРОЙКА СТРАНИЦЫ ---
st.set_page_config(
    page_title="Моя домашняя библиотека", page_icon="📚", layout="wide"
)

init_db()


# --- ФУНКЦИИ-КОЛБЭКИ ДЛЯ АВТОПЕРЕКЛЮЧЕНИЯ СТРАНИЦ ---
def cb_go_to_upload(book_id):
    st.session_state["reader_book_select_global"] = book_id
    st.session_state["sidebar_navigation"] = "📖 Читальный зал"

def cb_start_reading(book_id):
    st.session_state["global_reading_book_id"] = book_id
    st.session_state["global_start_time"] = datetime.now()
    st.session_state["reader_book_select_global"] = book_id
    st.session_state["sidebar_navigation"] = "📖 Читальный зал"

    # МАГИЯ: Автоматически переводим книгу в статус "Читаю" при старте таймера!
    conn = sqlite3.connect("library.db", timeout=10)
    cursor = conn.cursor()
    cursor.execute("UPDATE books SET status = 'Читаю' WHERE id = ?", (int(book_id),))
    conn.commit()
    conn.close()

def cb_change_book_type():
    # Извлекаем ID книги и выбранный тип из памяти виджета
    # Streamlit сам передает состояние виджета в session_state по его ключу
    for key in st.session_state.keys():
        if key.startswith("type_toggle_"):
            book_id = int(key.split("_")[-1])
            chosen_type = st.session_state[key]
            new_val = 1 if "Многотомное" in chosen_type else 0
            update_book_series_mode(book_id, new_val)

def cb_import_recommended_book(title, author):
    # Автоматически добавляем книгу в базу со статусом "Хочу прочитать"
    add_book(title, author, "Хочу прочитать", 0)
    st.session_state["sidebar_navigation"] = "📋 Моя библиотека"


# --- ИНТЕРФЕЙС ПРИЛОЖЕНИЯ ---
st.title("📚 Личный кабинет читателя")
st.write("Добро пожаловать в вашу цифровую библиотеку!")

# --- ГЛАВНОЕ НАВИГАЦИОННОЕ МЕНЮ В СИДБАРЕ ---
st.sidebar.title("🧭 Навигация")

menu_options = [
    "📋 Моя библиотека",
    "➕ Добавить книгу",
    "📊 Статистика",
    "⚙️ Настройки и Бонусы",
    "✍️ Любимые цитаты",
    "📖 Читальный зал",
    "🎯 Рекомендации книг",  # НОВЫЙ РАЗДЕЛ
]


# Инициализируем стартовую страницу в памяти ПРАВИЛЬНО
if "sidebar_navigation" not in st.session_state:
    st.session_state["sidebar_navigation"] = "📋 Моя библиотека"

# Создаем меню. Убрали параметр 'index' и привязали всё к 'key'
chosen_page = st.sidebar.radio(
    "Перейти в раздел:", options=menu_options, key="sidebar_navigation"
)

# Синхронизируем
st.session_state["current_page_name"] = st.session_state["sidebar_navigation"]

# --- СИСТЕМА ЗАЩИТЫ И АВТОРИЗАЦИИ ---
def check_password():
    """Возвращает True, если пользователь ввел правильный пароль."""
    if "password_correct" not in st.session_state:
        st.session_state["password_correct"] = False

    # Если пароль уже был успешно введен ранее в этой сессии, пропускаем дальше
    if st.session_state["password_correct"]:
        return True

    # Отображаем форму ввода пароля, если пользователь еще не авторизован
    st.subheader("🔒 Личный кабинет заблокирован")
    user_password = st.text_input("Введите ваш личный пароль для доступа к библиотеке:", type="password",
                                  key="auth_pwd_input")

    if st.button("🔑 Войти в библиотеку", use_container_width=True):
        if user_password == "LexaDEVLibrary133720":
            st.session_state["password_correct"] = True
            st.rerun()
        else:
            st.error("❌ Неверный пароль! Доступ к книгам запрещен.")

    return False


# Запускаем проверку. Если пароль не введен, принудительно останавливаем выполнение кода дальше!
if not check_password():
    st.stop()

current_tab = None

if "app_theme" not in st.session_state:
    st.session_state["app_theme"] = "default"

# Динамическое применение стилей CSS на лету
if st.session_state["app_theme"] == "forest":
    st.markdown(
        """
        <style>
        .stApp { background-color: #f0f7f4; }
        h1, h2, h3, p, span, label { color: #1e3f20 !important; }
        .stButton>button { background-color: #2d6a4f !important; color: white !important; border-radius: 8px; }
        .streamlit-expanderHeader { background-color: #d8f3dc !important; }
        </style>
    """,
        unsafe_allow_html=True,
    )

elif st.session_state["app_theme"] == "purple":
    st.markdown(
        """
        <style>
        .stApp { background-color: #1a1625; }
        h1, h2, h3, p, span, label { color: #e0aaff !important; }
        .stButton>button { background-color: #7b2cbf !important; color: white !important; border-radius: 8px; }
        .streamlit-expanderHeader { background-color: #3c096c !important; }
        div[data-baseweb="input"] { background-color: #240046 !important; }
        </style>
    """,
        unsafe_allow_html=True,
    )

class TabDummy:

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        pass


tab1, tab2, tab3, tab4, tab5, tab6, tab7 = (
    TabDummy(),
    TabDummy(),
    TabDummy(),
    TabDummy(),
    TabDummy(),
    TabDummy(),
    TabDummy(),
)

if st.session_state["current_page_name"] == "📋 Моя библиотека":
    current_tab = tab1
elif st.session_state["current_page_name"] == "➕ Добавить книгу":
    current_tab = tab2
elif st.session_state["current_page_name"] == "📊 Статистика":
    current_tab = tab3
elif st.session_state["current_page_name"] == "⚙️ Настройки и Бонусы":
    current_tab = tab4
elif st.session_state["current_page_name"] == "✍️ Любимые цитаты":
    current_tab = tab5
elif st.session_state["current_page_name"] == "📖 Читальный зал":
    current_tab = tab6
elif st.session_state["current_page_name"] == "🎯 Рекомендации книг":
    current_tab = tab7


# --- ВКЛАДКА 1: ПРОСМОТР И УПРАВЛЕНИЕ ---
if current_tab == tab1:
    books = get_books()
    if not books:
        st.info(
            "Ваша библиотека пока пуста. Перейдите во вкладку 'Добавить книгу'."
        )
    else:
        st.subheader("Список ваших книг")
        filter_col1, filter_col2 = st.columns(2)
        with filter_col1:
            search_query = st.text_input(
                "🔍 Поиск книги",
                placeholder="Введите название или автора для поиска...",
                key="search_input",
            )
        with filter_col2:
            status_filter = st.selectbox(
                "📌 Фильтр по статусу",
                ["Все статусы", "Хочу прочитать", "Читаю", "Прочитано"],
                key="status_filter_select",
            )

        filtered_books = books
        if search_query:
            filtered_books = [
                b
                for b in filtered_books
                if search_query.lower() in b[1].lower()
                or search_query.lower() in b[2].lower()
            ]
        if status_filter != "Все статусы":
            filtered_books = [
                b for b in filtered_books if b[3] == status_filter
            ]

        if filtered_books:
            df_export = pd.DataFrame(
                [
                    (b[1], b[2], b[3], b[4], b[5])
                    for b in filtered_books
                ],
                columns=[
                    "Название",
                    "Автор",
                    "Статус",
                    "Рейтинг",
                    "Дата добавления",
                ],
            )
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
                df_export.to_excel(writer, index=False, sheet_name="Книги")
            buffer.seek(0)
            st.download_button(
                label="📥 Скачать текущий список в Excel",
                data=buffer,
                file_name="my_library.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="dl_btn",
            )

        st.write("")
        if not filtered_books:
            st.warning("Книги с такими параметрами не найдены.")

        for book in filtered_books:
            b_id, b_title, b_author, b_status, b_rating, b_date = book
            stars = "⭐" * b_rating if b_rating > 0 else "Без оценки"
            with st.expander(f"📖 {b_title} — {b_author} ({stars})"):
                st.write(f"**Текущий статус:** {b_status}")
                st.write(
                    f"📅 **Дата добавления:** {b_date if b_date else 'Не указана'}"
                )
                col1, col2 = st.columns(2)
                with col1:
                    statuses = ["Хочу прочитать", "Читаю", "Прочитано"]
                    current_index = statuses.index(b_status)
                    new_status = st.selectbox(
                        "Сменить статус:",
                        statuses,
                        index=current_index,
                        key=f"status_{b_id}",
                    )
                    new_rating = st.slider(
                        "Изменить оценку:",
                        0,
                        5,
                        int(b_rating),
                        key=f"rate_{b_id}",
                    )
                    if new_status != b_status or new_rating != b_rating:
                        update_book_details(b_id, new_status, new_rating)
                        st.toast("Данные книги успешно обновлены!")
                        st.rerun()

                # --- БЛОК ЛИЧНЫХ КОММЕНТАРИЕВ И ЗАМЕТОК К КНИГЕ ---
                st.write("")
                st.write("📝 **Ваш читательский дневник / Комментарии:**")

                # Загружаем уже сохраненную ранее заметку из базы данных
                saved_notes = get_book_notes(b_id)

                # Поле ввода текста заметки
                new_notes = st.text_area(
                    "Запишите ваши мысли, инсайты или краткую рецензию:",
                    value=saved_notes,
                    key=f"notes_input_{b_id}",
                    placeholder="Напишите здесь что-нибудь о книге...",
                    max_chars=1000
                )

                # Кнопка сохранения комментария
                if st.button("💾 Сохранить заметку", key=f"save_notes_btn_{b_id}", use_container_width=True):
                    update_book_notes(b_id, new_notes)
                    st.toast("Ваш комментарий успешно сохранен!")
                    st.rerun()

                with col2:
                    st.write("**⏱️ Чтение книги**")

                    # 1. Узнаем формат книги (теперь вернет чистое 0 или 1)
                    is_series_mode = get_book_series_mode(b_id)

                    # 2. Умная проверка содержимого
                    has_content = False
                    if is_series_mode == 0:
                        # Если текст из базы не равен None и не пустой — значит книга загружена
                        book_text_content = get_book_text(b_id)
                        has_content = book_text_content is not None and book_text_content.strip() != ""
                        warning_msg = "📌 Сначала загрузите файл книги в 'Читальном зале'"
                        btn_label = "📥 Перейти к загрузке"
                    else:
                        has_content = len(get_book_volumes(b_id)) > 0
                        warning_msg = "📌 Сначала добавьте тома/части в 'Читальном зале'"
                        btn_label = "📥 Перейти к добавлению тома"

                    # 3. Логика отображения таймера
                    if not has_content:
                        st.warning(warning_msg)
                        if st.button(btn_label, key=f"go_to_upload_{b_id}", use_container_width=True,
                                     on_click=cb_go_to_upload, args=(b_id,)):
                            st.toast("Перенаправляем в Читальный зал...")
                            st.rerun()
                    else:
                        is_any_reading = "global_reading_book_id" in st.session_state
                        is_this_book_reading = is_any_reading and (st.session_state["global_reading_book_id"] == b_id)

                        if not is_any_reading:
                            if st.button("▶️ Начать читать", key=f"start_btn_{b_id}", use_container_width=True,
                                         on_click=cb_start_reading, args=(b_id,)):
                                st.toast(f"Открываем «{b_title}»...")
                                st.rerun()
                        else:
                            if is_this_book_reading:
                                st.info("⏳ Книга сейчас читается!")
                                if st.button("⏹️ Остановить чтение здесь", key=f"stop_here_btn_{b_id}", type="primary",
                                             use_container_width=True):
                                    elapsed_r = datetime.now() - st.session_state["global_start_time"]
                                    elapsed_min_r = max(1, int(elapsed_r.total_seconds() / 60))
                                    add_reading_session(b_id, elapsed_min_r)
                                    del st.session_state["global_reading_book_id"]
                                    del st.session_state["global_start_time"]
                                    st.success(f"Записано {elapsed_min_r} мин. в статистику!")
                                    st.rerun()
                            else:
                                st.button("▶️ Начать читать", key=f"start_btn_{b_id}", use_container_width=True,
                                          disabled=True)

                    st.divider()
                    if st.button("🗑️ Удалить книгу", key=f"del_{b_id}", type="primary"):
                        delete_book(b_id)
                        st.success("Книга удалена")
                        st.rerun()

# --- ВКЛАДКА 2: ДОБАВЛЕНИЕ КНИГИ ---
if current_tab == tab2:
    st.subheader("➕ Добавить новую книгу в библиотеку")
    with st.form("add_book_form_v3", clear_on_submit=True):
        title = st.text_input("Название книги:")
        author = st.text_input("Автор:")

        status = st.selectbox(
            "Статус:", ["Хочу прочитать", "Прочитано"]
        )

        submit = st.form_submit_button("Сохранить книгу!", key="final_save_btn")
        if submit:
            if title.strip() == "" or author.strip() == "":
                st.error("Пожалуйста, заполните название и автора!")
            else:
                # Передаем жесткий ноль (0) вместо переменной rating
                add_book(title, author, status, 0)
                st.success(f"Книга «{title}» успешно добавлена!")
                st.rerun()

# --- ВКЛАДКА 3: СТАТИСТИКА ---
if current_tab == tab3:
    st.subheader("📊 Аналитика вашей библиотеки")
    books = get_books()
    if not books:
        st.info("Нет данных для отображения статистики.")
    else:
        total_books = len(books)
        statuses_list = [b[3] for b in books]
        ratings_list = [b[4] for b in books if b[4] > 0]
        avg_rating = (
            round(sum(ratings_list) / len(ratings_list), 1)
            if ratings_list
            else 0.0
        )

        total_minutes = get_total_reading_time()

        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Всего книг в базе", total_books)
        m2.metric("Прочитано книг ✨", statuses_list.count("Прочитано"))
        m3.metric(
            "Средняя оценка",
            f"{avg_rating} / 5" if avg_rating else "Нет оценок",
        )
        m4.metric("Время чтения ⏱️", f"{total_minutes} мин.")

        st.divider()
        count_read = statuses_list.count("Прочитано")
        if total_books > 0:
            progress_percent = int((count_read / total_books) * 100)
            st.write(
                f"🏆 **Прогресс чтения:** прочитано **{progress_percent}%** от общего фонда библиотеки"
            )
            st.progress(progress_percent / 100)

        st.write("")
        st.subheader("📈 Книги по статусам")
        chart_data = pd.DataFrame(
            {
                "Статус": ["Хочу прочитать", "Читаю", "Прочитано"],
                "Количество": [
                    statuses_list.count("Хочу прочитать"),
                    statuses_list.count("Читаю"),
                    statuses_list.count("Прочитано"),
                ],
            }
        )
        st.bar_chart(
            data=chart_data,
            x="Статус",
            y="Количество",
            color="Статус",
            use_container_width=True,
        )

        st.subheader("📅 Динамика добавления книг")
        months = [b[5][:7] for b in books if b[5]]
        if months:
            df_months = (
                pd.DataFrame(months, columns=["Месяц"])
                .value_counts()
                .reset_index()
            )
            df_months.columns = ["Месяц", "Книг добавлено"]
            df_months = df_months.sort_values(by="Месяц")
            st.line_chart(
                data=df_months,
                x="Месяц",
                y="Книг добавлено",
                use_container_width=True,
            )
        else:
            st.info("Нет данных о датах для отображения динамики.")

# --- ВКЛАДКА 4: НАСТРОЙКИ И БОНУСЫ ---
if current_tab == tab4:
    st.subheader("🏆 Личный статус читателя")

    # Расчитываем все игровые показатели на лету
    status = get_reader_status_data()

    # Создаем красивую карточку профиля с золотыми звездами
    star_rating = "⭐" * status["stars"] + "🔲" * (5 - status["stars"])

    col_prof1, col_prof2 = st.columns([1, 2])
    with col_prof1:
        st.markdown(f"### Уровень {status['level']}")
        st.write(f"**Ранг:** {status['rank']}")
        st.write(f"**Рейтинг:** {star_rating}")
    with col_prof2:
        st.info(f"📜 *{status['desc']}*")
        st.write(f"📊 **Всего начитано:** {status['minutes']} мин. ({status['xp']} XP)")

    # Шкала прогресса до следующего уровня
    st.caption(f"Прогресс уровня: {status['current_level_xp']} / {status['next_level_xp']} XP")
    st.progress(status["progress"] / 100)

    st.divider()  # Отделяем статус от Книжного компаса

    st.subheader("🎲 Книжный компас (Что почитать?)")
    st.write("Позвольте алгоритму случайного выбора определить вашу следующую книгу из списка **«Хочу прочитать»**!")

    # Кнопка генератора случайной книги
    if st.button(
            "🎰 Выбрать случайную книгу",
            type="secondary",
            use_container_width=True,
            key="rnd_btn",
    ):
        random_book = get_random_wishlist_book()
        if random_book:
            title, author = random_book
            st.balloons()  # Запуск праздничной анимации шаров
            st.success(f"🎉 Ваша следующая книга: **«{title}»** — {author}")
        else:
            st.info(
                "💡 В вашем списке 'Хочу прочитать' пока нет книг. Добавьте их во 2-й вкладке!"
            )

    st.divider()

    st.subheader("🎨 Кастомные темы оформления")
    st.write("Смените цветовую палитру интерфейса одним кликом:")

    # Инициализируем переменную темы в памяти приложения, если её там нет
    if "app_theme" not in st.session_state:
        st.session_state["app_theme"] = "default"

    # Кнопки выбора тем оформления в один ряд
    t_col1, t_col2, t_col3 = st.columns(3)
    with t_col1:
        if st.button("🌲 Лесная прохлада", use_container_width=True, key="t1"):
            st.session_state["app_theme"] = "forest"
            st.rerun()
    with t_col2:
        if st.button("🍇 Ночной аметист", use_container_width=True, key="t2"):
            st.session_state["app_theme"] = "purple"
            st.rerun()
    with t_col3:
        if st.button("⚪ Сбросить стиль", use_container_width=True, key="t3"):
            st.session_state["app_theme"] = "default"
            st.rerun()

    st.write("")  # Небольшой пустой отступ для красоты
    if st.session_state["app_theme"] == "forest":
        st.caption("Применена тема: *Лесная прохлада*")
    elif st.session_state["app_theme"] == "purple":
        st.caption("Применена тема: *Ночной аметист*")
    else:
        st.caption("Применена тема: *Стандартная*")

# --- ВКЛАДКА 5: ЛЮБИМЫЕ ЦИТАТЫ ---
if current_tab == tab5:
    st.subheader("✍️ Ваша галерея мудрых мыслей")

    all_books = get_books()

    if not all_books:
        st.info(
            "Сначала добавьте хотя бы одну книгу в библиотеку, чтобы сохранять цитаты."
        )
    else:
        # Форма добавления новой цитаты
        with st.form("add_quote_form", clear_on_submit=True):
            st.write("**Добавить новую цитату:**")

            # Выпадающий список книг для привязки цитаты
            book_options = {b[0]: f"«{b[1]}» — {b[2]}" for b in all_books}
            selected_book_id = st.selectbox(
                "Выберите книгу:",
                options=list(book_options.keys()),
                format_func=lambda x: book_options[x],
                key="quote_book_select",
            )

            quote_text = st.text_area(
                "Текст цитаты:", placeholder="Введите красивую фразу..."
            )
            page_num = st.text_input(
                "Страница (необязательно):", placeholder="Например: 142"
            )

            submit_quote = st.form_submit_button(
                "Сохранить цитату", key="save_quote_btn"
            )

            if submit_quote:
                if quote_text.strip() == "":
                    st.error("Текст цитаты не может быть пустым!")
                else:
                    add_quote(selected_book_id, quote_text, page_num)
                    st.success("Цитата успешно сохранена!")
                    st.rerun()

        st.divider()

        # Отображение сохраненных цитат
        st.subheader("📌 Сохраненные цитаты")
        quotes = get_all_quotes()

        if not quotes:
            st.caption("Вы пока не сохранили ни одной цитаты.")
        else:
            for q in quotes:
                q_id, b_title, b_author, q_text, q_page, q_date = q
                page_str = f" (стр. {q_page})" if q_page else ""

                # Оформляем цитату в красивый блок
                st.markdown(
                    f"""
                    <div style="background-color: #f9f9f9; padding: 15px; border-left: 5px solid #7b2cbf; border-radius: 5px; margin-bottom: 10px;">
                        <p style="font-style: italic; color: #333; font-size: 16px;">"{q_text}"</p>
                        <small style="color: #666;">📖 <b>{b_title}</b> — {b_author}{page_str} | 🗓️ {q_date}</small>
                    </div>
                """,
                    unsafe_allow_html=True,
                )

                # Кнопка удаления цитаты
                if st.button(
                    "🗑️ Удалить цитату", key=f"del_q_{q_id}", type="secondary"
                ):
                    delete_quote(q_id)
                    st.success("Цитата удалена")
                    st.rerun()

# --- ВКЛАДКА 6: ЧИТАЛЬНЫЙ ЗАЛ ---
if current_tab == tab6:
    st.subheader("📖 Электронный читальный зал")

    all_books = get_books()

    if not all_books:
        st.info("Сначала добавьте книгу в библиотеку, чтобы открыть её для чтения.")
    else:
        book_options = {b[0]: f"«{b[1]}» — {b[2]}" for b in all_books}

        fallback_book_id = list(book_options.keys())[0] if book_options else None
        default_book = st.session_state.get("global_reading_book_id", fallback_book_id)

        if default_book not in book_options:
            default_book = fallback_book_id

        default_index = list(book_options.keys()).index(default_book) if default_book in book_options else 0

        selected_read_id = st.selectbox(
            "Какую книгу будем читать?",
            options=list(book_options.keys()),
            format_func=lambda x: book_options[x],
            index=default_index,
            key="reader_book_select_global",
        )

        is_series_mode = get_book_series_mode(selected_read_id)

        st.write("")
        book_type = st.radio(
            "Формат этого издания:",
            ["Одиночная книга (один файл)", "Многотомное издание (несколько частей/томов)"],
            index=1 if is_series_mode == 1 else 0,
            key=f"type_toggle_{selected_read_id}",
            on_change=cb_change_book_type
        )

        st.divider()

        new_mode_value = 1 if "Многотомное" in book_type else 0

        # ==========================================
        # РЕЖИМ 1: ОДИНОЧНАЯ КНИГА
        # ==========================================
        if new_mode_value == 0:
            saved_text = get_book_text(selected_read_id)

            if saved_text is None:
                st.write("📌 Для этой книги еще не загружен электронный текст.")
                uploaded_file = st.file_uploader(
                    "Загрузите текст книги (.txt, .html, .fb2, .epub, .docx, .pdf)",
                    type=["txt", "html", "fb2", "epub", "docx", "pdf"],
                    key=f"single_file_{selected_read_id}"
                )

                if uploaded_file is not None:
                    file_type = uploaded_file.name.split(".")[-1].lower()
                    raw_bytes = uploaded_file.getvalue()
                    uploaded_text = ""

                    # --- 1. Формат FB2 ---
                    if file_type == "fb2":
                        import xml.etree.ElementTree as ET

                        try:
                            root = ET.fromstring(raw_bytes)
                            paragraphs_fb2 = []
                            for elem in root.iter():
                                if elem.tag.endswith('p') and elem.text:
                                    paragraphs_fb2.append(elem.text.strip())
                            uploaded_text = "\n\n".join(paragraphs_fb2)
                        except Exception as e:
                            st.error(f"Ошибка чтения FB2: {e}")

                    # --- 2. Формат EPUB ---
                    elif file_type == "epub":
                        import ebooklib
                        from ebooklib import epub
                        import io
                        from bs4 import BeautifulSoup

                        try:
                            epub_file = io.BytesIO(raw_bytes)
                            book = epub.read_epub(epub_file)
                            epub_paragraphs = []
                            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                                soup = BeautifulSoup(item.get_content(), 'html.parser')
                                text_content = soup.get_text()
                                if text_content.strip():
                                    epub_paragraphs.append(text_content.strip())
                            uploaded_text = "\n\n".join(epub_paragraphs)
                        except Exception as e:
                            st.error(f"Ошибка чтения EPUB: {e}")

                    # --- 3. Формат DOCX ---
                    elif file_type == "docx":
                        import docx
                        import io

                        try:
                            docx_file = io.BytesIO(raw_bytes)
                            doc = docx.Document(docx_file)
                            docx_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                            uploaded_text = "\n\n".join(docx_paragraphs)
                        except Exception as e:
                            st.error(f"Ошибка чтения DOCX: {e}")

                    # --- 4. Формат PDF ---
                    elif file_type == "pdf":
                        import pypdf
                        import io

                        try:
                            pdf_file = io.BytesIO(raw_bytes)
                            reader = pypdf.PdfReader(pdf_file)
                            pdf_pages = []
                            for page in reader.pages:
                                text = page.extract_text()
                                if text:
                                    pdf_pages.append(text.strip())
                            uploaded_text = "\n\n".join(pdf_pages)
                        except Exception as e:
                            st.error(f"Ошибка чтения PDF: {e}")

                    # --- 5. Текстовые TXT и HTML ---
                    else:
                        for encoding in ["utf-8", "windows-1251", "utf-16"]:
                            try:
                                uploaded_text = raw_bytes.decode(encoding)
                                break
                            except UnicodeDecodeError:
                                continue

                        if file_type == "html" and uploaded_text:
                            import re

                            uploaded_text = re.sub(r'<script.*?</script>', '', uploaded_text, flags=re.DOTALL)
                            uploaded_text = re.sub(r'<style.*?</style>', '', uploaded_text, flags=re.DOTALL)
                            uploaded_text = re.sub(r'</p>', '\n\n', uploaded_text)
                            uploaded_text = re.sub(r'<br\s*/?>', '\n\n', uploaded_text)
                            uploaded_text = re.sub(r'<.*?>', '', uploaded_text)

                    if uploaded_text.strip():
                        save_book_text(selected_read_id, uploaded_text)
                        st.success("🎉 Текст книги успешно сохранен!")
                        st.rerun()
                    else:
                        st.error("Не удалось извлечь текст из файла.")
            else:
                raw_paragraphs = [p.strip() for p in saved_text.split("\n\n") if p.strip()]
                paragraphs = []
                chunk = []
                for p in raw_paragraphs:
                    chunk.append(p)
                    if len(chunk) >= 4 or len(" ".join(chunk)) > 1500:
                        paragraphs.append("\n\n".join(chunk))
                        chunk = []
                if chunk:
                    paragraphs.append("\n\n".join(chunk))

                if not paragraphs:
                    st.warning("Книга пуста.")
                else:
                    total_pages = len(paragraphs)
                    saved_pos = get_book_position(selected_read_id)
                    current_page = max(0, min(saved_pos, total_pages - 1))

                    is_reading_now = "global_reading_book_id" in st.session_state and st.session_state[
                        "global_reading_book_id"] == selected_read_id
                    t_col1, t_col2 = st.columns(2)
                    with t_col1:
                        if not is_reading_now:
                            if st.button("▶️ Начать сессию чтения", type="secondary", use_container_width=True,
                                         key="single_start"):
                                st.session_state["global_reading_book_id"] = selected_read_id
                                st.session_state["global_start_time"] = datetime.now()
                                st.rerun()
                        else:
                            elapsed_r = datetime.now() - st.session_state["global_start_time"]
                            elapsed_min_r = max(1, int(elapsed_r.total_seconds() / 60))
                            st.write(f"⏳ Читаете: {int(elapsed_r.total_seconds())} сек.")
                            if st.button("⏹️ Завершить и сохранить время", type="primary", use_container_width=True,
                                         key="single_stop"):
                                add_reading_session(selected_read_id, elapsed_min_r)
                                del st.session_state["global_reading_book_id"]
                                del st.session_state["global_start_time"]
                                st.success(f"Записано {elapsed_min_r} мин.!")
                                st.rerun()
                    with t_col2:
                        st.caption(f"📊 Прогресс: страница {current_page + 1} из {total_pages}")
                        st.progress((current_page + 1) / total_pages)

                    st.write("")
                    st.markdown(
                        f"""
                        <div style="background-color: #fffdf6; padding: 25px; border: 1px solid #e6dfb8; border-radius: 8px; font-size: 18px; line-height: 1.6; color: #2c2512; font-family: 'Georgia', serif; min-height: 200px; margin-bottom: 20px;">
                            {paragraphs[current_page].replace('\n\n', '<br><br>')}
                        </div>
                    """,
                        unsafe_allow_html=True
                    )

                    nav_col1, nav_col2, nav_col3 = st.columns(3)
                    with nav_col1:
                        if st.button("⬅️ Назад", use_container_width=True, disabled=(current_page == 0), key="s_back"):
                            current_page -= 1
                            update_book_position(selected_read_id, current_page)
                            st.rerun()
                    with nav_col2:
                        if st.button("🗑️ Удалить текст книги", use_container_width=True, type="primary", key="s_del"):
                            save_book_text(selected_read_id, None)
                            update_book_position(selected_read_id, 0)
                            st.rerun()
                    with nav_col3:
                        if st.button("Вперед ➡️", use_container_width=True, disabled=(current_page == total_pages - 1), key="s_next"):
                            current_page += 1
                            update_book_position(selected_read_id, current_page)
                            st.rerun()

        # ==========================================
        # РЕЖИМ 2: МНОГОТОМНОЕ ИЗДАНИЕ
        # ==========================================
        else:
            volumes = get_book_volumes(selected_read_id)

            with st.expander("➕ Добавить новый том / часть к этой книге"):
                with st.form(f"add_volume_form_{selected_read_id}", clear_on_submit=True):
                    vol_title = st.text_input("Название тома/части:",
                                              placeholder="Например: Том 1. Братство кольца")
                    uploaded_file = st.file_uploader(
                        "Загрузите файл тома (.txt, .html, .fb2, .epub, .docx, .pdf):",
                        type=["txt", "html", "fb2", "epub", "docx", "pdf"])
                    submit_vol = st.form_submit_button("📁 Загрузить и сохранить часть")

                    if submit_vol:
                        if vol_title.strip() == "":
                            st.error("Пожалуйста, введите название тома!")
                        elif uploaded_file is None:
                            st.error("Пожалуйста, выберите файл книги!")
                        else:
                            file_type = uploaded_file.name.split(".")[-1].lower()
                            raw_bytes = uploaded_file.getvalue()
                            uploaded_text = ""

                            # --- 1. Формат FB2 ---
                            if file_type == "fb2":
                                import xml.etree.ElementTree as ET

                                try:
                                    root = ET.fromstring(raw_bytes)
                                    paragraphs_fb2 = []
                                    for elem in root.iter():
                                        if elem.tag.endswith('p') and elem.text:
                                            paragraphs_fb2.append(elem.text.strip())
                                    uploaded_text = "\n\n".join(paragraphs_fb2)
                                except Exception as e:
                                    st.error(f"Ошибка чтения FB2: {e}")

                            # --- 2. Формат EPUB ---
                            elif file_type == "epub":
                                import ebooklib
                                from ebooklib import epub
                                import io
                                from bs4 import BeautifulSoup

                                try:
                                    epub_file = io.BytesIO(raw_bytes)
                                    book = epub.read_epub(epub_file)
                                    epub_paragraphs = []
                                    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                                        soup = BeautifulSoup(item.get_content(), 'html.parser')
                                        text_content = soup.get_text()
                                        if text_content.strip():
                                            epub_paragraphs.append(text_content.strip())
                                    uploaded_text = "\n\n".join(epub_paragraphs)
                                except Exception as e:
                                    st.error(f"Ошибка чтения EPUB: {e}")

                            # --- 3. Формат DOCX ---
                            elif file_type == "docx":
                                import docx
                                import io

                                try:
                                    docx_file = io.BytesIO(raw_bytes)
                                    doc = docx.Document(docx_file)
                                    docx_paragraphs = [p.text.strip() for p in doc.paragraphs if
                                                       p.text.strip()]
                                    uploaded_text = "\n\n".join(docx_paragraphs)
                                except Exception as e:
                                    st.error(f"Ошибка чтения DOCX: {e}")

                            # --- 4. Формат PDF ---
                            elif file_type == "pdf":
                                import pypdf
                                import io

                                try:
                                    pdf_file = io.BytesIO(raw_bytes)
                                    reader = pypdf.PdfReader(pdf_file)
                                    pdf_pages = []
                                    for page in reader.pages:
                                        text = page.extract_text()
                                        if text:
                                            pdf_pages.append(text.strip())
                                    uploaded_text = "\n\n".join(pdf_pages)
                                except Exception as e:
                                    st.error(f"Ошибка чтения PDF: {e}")

                            # --- 5. Текстовые файлы ---
                            else:
                                for encoding in ["utf-8", "windows-1251", "utf-16"]:
                                    try:
                                        uploaded_text = raw_bytes.decode(encoding)
                                        break
                                    except UnicodeDecodeError:
                                        continue

                                if file_type == "html" and uploaded_text:
                                    import re

                                    uploaded_text = re.sub(r'<script.*?</script>', '', uploaded_text,
                                                           flags=re.DOTALL)
                                    uploaded_text = re.sub(r'<style.*?</style>', '', uploaded_text,
                                                           flags=re.DOTALL)
                                    uploaded_text = re.sub(r'</p>', '\n\n', uploaded_text)
                                    uploaded_text = re.sub(r'<br\s*/?>', '\n\n', uploaded_text)
                                    uploaded_text = re.sub(r'<.*?>', '', uploaded_text)

                            if uploaded_text.strip():
                                add_book_volume(selected_read_id, vol_title, uploaded_text)
                                st.success(f"🎉 «{vol_title}» успешно добавлен!")
                                st.rerun()
                            else:
                                st.error("Не удалось извлечь текст из файла.")

            if not volumes:
                st.info("💡 У этой книги пока нет загруженных томов. Добавьте первый том выше!")
            else:
                # Превращаем ID тома в строку для безопасности Streamlit
                vol_options = {str(v[0]): v[1] for v in volumes}

                selected_vol_key = st.selectbox(
                    "Выберите том/часть для чтения:",
                    options=list(vol_options.keys()),
                    format_func=lambda x: vol_options[x],
                    key=f"vol_select_{selected_read_id}"
                )

                # Обратно превращаем в число для корректного SQL-запроса в базу данных
                selected_vol_id = int(selected_vol_key)

                saved_text, saved_pos = get_volume_data(selected_vol_id)

                if not saved_text:
                    st.warning("Текст этого тома пуст.")
                else:
                    raw_paragraphs = [p.strip() for p in saved_text.split("\n\n") if p.strip()]
                    paragraphs = []
                    chunk = []
                    for p in raw_paragraphs:
                        chunk.append(p)
                        if len(chunk) >= 4 or len(" ".join(chunk)) > 1500:
                            paragraphs.append("\n\n".join(chunk))
                            chunk = []
                    if chunk:
                        paragraphs.append("\n\n".join(chunk))

                    if not paragraphs:
                        st.warning("Ошибка обработки текста.")
                    else:
                        total_pages = len(paragraphs)
                        current_page = max(0, min(saved_pos, total_pages - 1))

                        is_reading_now = "global_reading_book_id" in st.session_state and st.session_state[
                            "global_reading_book_id"] == selected_read_id
                        t_col1, t_col2 = st.columns(2)
                        with t_col1:
                            if not is_reading_now:
                                if st.button("▶️ Начать сессию чтения", type="secondary", use_container_width=True,
                                             key="m_start"):
                                    st.session_state["global_reading_book_id"] = selected_read_id
                                    st.session_state["global_start_time"] = datetime.now()
                                    st.rerun()
                            else:
                                elapsed_r = datetime.now() - st.session_state["global_start_time"]
                                elapsed_min_r = max(1, int(elapsed_r.total_seconds() / 60))
                                st.write(f"⏳ Сессия активна! Читаете: {int(elapsed_r.total_seconds())} сек.")
                                if st.button("⏹️ Завершить и сохранить время", type="primary", use_container_width=True,
                                             key="m_stop"):
                                    add_reading_session(selected_read_id, elapsed_min_r)
                                    del st.session_state["global_reading_book_id"]
                                    del st.session_state["global_start_time"]
                                    st.success(f"Записано {elapsed_min_r} мин.!")
                                    st.rerun()
                        with t_col2:
                            st.caption(f"📊 Прогресс тома: страница {current_page + 1} из {total_pages}")
                            st.progress((current_page + 1) / total_pages)

                        st.write("")
                        st.markdown(
                            f"""
                            <div style="background-color: #fffdf6; padding: 25px; border: 1px solid #e6dfb8; border-radius: 8px; font-size: 18px; line-height: 1.6; color: #2c2512; font-family: 'Georgia', serif; min-height: 200px; margin-bottom: 20px;">
                                {paragraphs[current_page].replace('\n\n', '<br><br>')}
                            </div>
                        """,
                            unsafe_allow_html=True
                        )

                        nav_col1, nav_col2, nav_col3 = st.columns(3)
                        with nav_col1:
                            if st.button("⬅️ Назад", use_container_width=True, disabled=(current_page == 0),
                                         key="m_back"):
                                current_page -= 1
                                update_volume_position(selected_vol_id, current_page)
                                st.rerun()
                        with nav_col2:
                            if st.button("🗑️ Удалить текст этого тома", use_container_width=True, type="primary", key="m_del"):
                                delete_book_volume(selected_vol_id)
                                st.rerun()
                        with nav_col3:
                            if st.button("Вперед ➡️", use_container_width=True,
                                         disabled=(current_page == total_pages - 1), key="m_next"):
                                current_page += 1
                                update_volume_position(selected_vol_id, current_page)
                                st.rerun()

# --- ВКЛАДКА 7: РЕКОМЕНДАЦИИ КНИГ ---
if current_tab == tab7:
    st.subheader("🎯 Книжная полка рекомендаций")
    st.write(
        "Специально подобранные шедевры литературы. Нажмите кнопку, чтобы мгновенно добавить книгу в свой список к прочтению!"
    )

    # База данных рекомендаций (Название, Автор, Описание)
    recommendations = {
        "🛸 Научная фантастика": [
            (
                "Дюна",
                "Фрэнк Герберт",
                "Эпическая сага о песчаной планете Арракис, интригах, пророках и борьбе за самый ценный ресурс во Вселенной.",
            ),
            (
                "Основание",
                "Айзек Азимов",
                "История о падении Галактической Империи и ученых, которые пытаются сохранить знания человечества.",
            ),
        ],
        "🧠 Саморазвитие и Психология": [
            (
                "Атомные привычки",
                "Джеймс Клир",
                "Простая и понятная книга о том, как крошечные изменения в поведении приводят к колоссальным результатам в жизни.",
            ),
            (
                "Думай медленно... решай быстро",
                "Даниэль Канеман",
                "Шедевр от нобелевского лауреата о двух системах ума, которые управляют нашими решениями и ошибками.",
            ),
        ],
        "🏛️ Мировая Классика": [
            (
                "1984",
                "Джордж Оруэлл",
                "Главный антиутопический роман века о Большом Брате, тотальном контроле, любви и свободе личности.",
            ),
            (
                "Мастер и Маргарита",
                "Михаил Булгаков",
                "Бессмертный мистический роман о визите дьявола в Москву, великой любви и древнем Понтии Пилате.",
            ),
        ],
    }

    # Отображаем жанры в виде красивых раскрывающихся блоков
    for genre, books_list in recommendations.items():
        st.write("")
        st.markdown(f"### {genre}")

        for r_title, r_author, r_desc in books_list:
            # Оформляем каждую книгу в виде стильного блока
            with st.container(border=True):
                col_rec1, col_rec2 = st.columns([3, 1])

                with col_rec1:
                    st.markdown(f"**«{r_title}»** — *{r_author}*")
                    st.caption(r_desc)

                with col_rec2:
                    st.write("")  # Отступ для выравнивания кнопки
                    # Кнопка импорта с колбэком
                    st.button(
                        "📥 В список",
                        key=f"import_{r_title.replace(' ', '_')}",
                        use_container_width=True,
                        type="secondary",
                        on_click=cb_import_recommended_book,
                        args=(r_title, r_author),
                    )
